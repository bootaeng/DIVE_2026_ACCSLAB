"""Extract text from .docx files using python-docx."""
import os
import sys
import io

# Force UTF-8 output
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

OUTPUT_FILE = os.path.join(BASE_DIR, "code", "extracted_text.txt")

files = [
    os.path.join(BASE_DIR, "회의록.docx"),
    os.path.join(BASE_DIR, "부산교통공사x짐캐리_해커톤_기획보고서.docx"),
]

with open(OUTPUT_FILE, "w", encoding="utf-8") as out:
    for filepath in files:
        filename = os.path.basename(filepath)
        out.write("=" * 80 + "\n")
        out.write(f"FILE: {filename}\n")
        out.write("=" * 80 + "\n")
        doc = Document(filepath)

        # Extract paragraphs
        for para in doc.paragraphs:
            out.write(para.text + "\n")

        # Extract text from tables
        for table_idx, table in enumerate(doc.tables):
            out.write(f"\n[TABLE {table_idx + 1}]\n")
            for row in table.rows:
                row_text = "\t".join(cell.text for cell in row.cells)
                out.write(row_text + "\n")

        out.write("\n\n")

print(f"Extracted text saved to: {OUTPUT_FILE}")
