import os
import sys
import io
import glob
import fitz  # PyMuPDF
from docx import Document

# Force UTF-8 output
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUTPUT_FILE = os.path.join(BASE_DIR, "code", "extracted_text.txt")

def extract_docx(filepath):
    text_content = []
    try:
        doc = Document(filepath)
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_content.append(para.text)
        
        # Extract text from tables
        for table_idx, table in enumerate(doc.tables):
            text_content.append(f"\n[TABLE {table_idx + 1}]")
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_content.append(row_text)
    except Exception as e:
        text_content.append(f"ERROR reading docx {os.path.basename(filepath)}: {e}")
    return "\n".join(text_content)

def extract_pdf(filepath):
    text_content = []
    try:
        doc = fitz.open(filepath)
        total_pages = len(doc)
        text_content.append(f"(Total pages: {total_pages})\n")
        
        for page_num in range(total_pages):
            page = doc[page_num]
            text = page.get_text()
            text_content.append(f"--- Page {page_num + 1} ---")
            if text.strip():
                text_content.append(text.strip())
            else:
                text_content.append("[No extractable text on this page - likely image-based]")
        doc.close()
    except Exception as e:
        text_content.append(f"ERROR reading pdf {os.path.basename(filepath)}: {e}")
    return "\n".join(text_content)

def main():
    print("Starting unified text extraction...")
    
    # Get all .docx and .pdf files in BASE_DIR
    docx_files = sorted(glob.glob(os.path.join(BASE_DIR, "*.docx")))
    pdf_files = sorted(glob.glob(os.path.join(BASE_DIR, "*.pdf")))
    
    print(f"Found {len(docx_files)} Word files and {len(pdf_files)} PDF files.")
    
    with open(OUTPUT_FILE, "w", encoding="utf-8") as out:
        # Process Word files
        for filepath in docx_files:
            filename = os.path.basename(filepath)
            separator = "=" * 80
            header = f"{separator}\nFILE: {filename}\n{separator}\n"
            print(f"Extracting: {filename}")
            out.write(header)
            
            content = extract_docx(filepath)
            out.write(content + "\n\n")
            
        # Process PDF files
        for filepath in pdf_files:
            filename = os.path.basename(filepath)
            separator = "=" * 80
            header = f"{separator}\nFILE: {filename}\n{separator}\n"
            print(f"Extracting: {filename}")
            out.write(header)
            
            content = extract_pdf(filepath)
            out.write(content + "\n\n")
            
    print(f"\nAll extracted text successfully saved to:\n{OUTPUT_FILE}")

if __name__ == "__main__":
    main()
